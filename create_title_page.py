
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(run, font_name, font_size, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold

def create_title_page():
    doc = Document()
    
    # Set margins to 2cm as per guidelines
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Helper to add empty lines
    def add_empty_lines(count):
        for _ in range(count):
            doc.add_paragraph()

    # --- Header / Top Section ---
    # "EXAMINATION OFFICE" / "IU.ORG" based on the guideline visual
    # usually top left or right, let's put it top left as plain text for now or standard layout
    # Guidelines text: "page 1 of 5 EXAMINATION OFFICE IU.ORG" seems to be header content.
    # We will simulate a clean professional header.
    
    # p = doc.add_paragraph()
    # run = p.add_run("EXAMINATION OFFICE\nIU.ORG")
    # set_font(run, "Arial", 11)
    # p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    add_empty_lines(15) # Adjust spacing to push title to center-ish
    
    # --- Thesis Type ---
    p = doc.add_paragraph()
    run = p.add_run("MASTER THESIS")
    set_font(run, "Arial", 16, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_empty_lines(2)
    
    # --- Title ---
    title_text = "Empirical Analysis of Accuracy–Privacy Trade-offs in Federated Learning for Mobile Movie Recommendation Systems"
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    set_font(run, "Arial", 20, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_empty_lines(8)
    
    # --- Course Details ---
    # Usually in a block
    
    details = [
        ("Course of Study:", "M.Sc. in Artificial Intelligence (120 ECTS)"),
        ("Course Name:", "Master Thesis"), # Using 'Master Thesis' as course name if not specified otherwise, but user said 'M.Sc...' is Course Name? User input: "Course Name is M.Sc. in Artificial Intelligence (120 ECTS)". Wait, usually Course Name is the module (e.g. 'Master Thesis') and Course of Study is the degree. User put the same? Re-reading: "Course Name is M.Sc. in Artificial Intelligence". Guidelines say: "course name, course of study". I will put specific inputs.
        # User input:
        # Course Name: M.Sc. in Artificial Intelligence (120 ECTS)
        # Author: Jonathan John
        # Matriculation: 92126351
        # Tutor: Duong-Trung, Nghia, Dr.
        # Date: 2026-02-18
    ]
    
    # Using a table for alignment is often cleaner for the details section
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False 
    # Set column widths? 
    # Let's just use paragraphs for simplicity and standard look
    
    def add_detail_line(label, value):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label} ")
        set_font(run_label, "Arial", 12, bold=True)
        if value:
            run_val = p.add_run(value)
            set_font(run_val, "Arial", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Course Name according to user input
    add_detail_line("Course of Study:", "M.Sc. in Artificial Intelligence (120 ECTS)")
    # add_detail_line("Course Name:", "Master Thesis") # Often explicit? Let's assume standard if not provided separately or implicit. User said "Course Name IS M.Sc...". I will follow user input strictly?
    # Actually, guidelines say "course name, course of study". User gave "Course Name is M.Sc...". I will label it 'Course:'
    
    # add_detail_line("Course:", "M.Sc. in Artificial Intelligence (120 ECTS)") 
    # Wait, usually it is:
    # Course of Study: M.Sc. Artificial Intelligence
    # Course: Master Thesis (or similar module name)
    # User input might have combined them. I will list what was provided.
    
    # Let's try to infer standard structure:
    # Author
    add_detail_line("Author:", "Jonathan John")
    add_detail_line("Matriculation Number:", "92126351")
    
    add_empty_lines(1)
    
    add_detail_line("Tutor:", "Duong-Trung, Nghia, Dr.")
    
    add_empty_lines(1)
    
    add_detail_line("Date:", "2026-02-18")

    doc.save("Title_Page.docx")
    print("Title_Page.docx created successfully.")

if __name__ == "__main__":
    create_title_page()
