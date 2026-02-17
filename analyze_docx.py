import docx
from docx.document import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within parent, in document order.
    Each returned value is an instance of either Table or Paragraph.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)

def analyze(doc_path):
    doc = docx.Document(doc_path)
    
    print(f"Analyzing {doc_path}...")
    
    # Iterate through paragraphs to find images
    # Images in python-docx are usually found via runs or blips in the XML
    
    count = 0
    for i, para in enumerate(doc.paragraphs):
        # Naive check for images in runs
        has_image = False
        for run in para.runs:
            if 'w:drawing' in run._element.xml:
                has_image = True
                break
        
        if has_image:
            count += 1
            print(f"\n[Image found at Paragraph {i}]")
            # Print surrounding context
            prev_text = doc.paragraphs[i-1].text if i > 0 else "[Start of Doc]"
            next_text = doc.paragraphs[i+1].text if i < len(doc.paragraphs)-1 else "[End of Doc]"
            print(f"  Prev: {prev_text[:100]}...")
            print(f"  Curr (Image Para): {para.text[:100]}...")
            print(f"  Next: {next_text[:100]}...")

if __name__ == "__main__":
    analyze("Empirical Analysis of Accuracy.docx")
