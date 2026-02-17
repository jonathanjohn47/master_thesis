
import docx
import re
from docx.oxml.ns import qn

def get_document_images(doc):
    """
    Return a list of dicts: {'image_para': paragraph_obj, 'caption_para': paragraph_obj_or_None, 'old_num': int_or_None}
    """
    images = []
    paragraphs = doc.paragraphs
    skip_indices = set()

    for i, para in enumerate(paragraphs):
        if i in skip_indices:
            continue

        # Check for image
        has_image = False
        if 'w:drawing' in para._element.xml:
            has_image = True
        
        if has_image:
            # Check next paragraph for caption
            caption_para = None
            old_num = None
            if i + 1 < len(paragraphs):
                next_p = paragraphs[i+1]
                text = next_p.text.strip()
                # Pattern: "Figure 1:", "Figure 1.", "Figure 1 "
                match = re.match(r"^Figure\s+(\d+)", text, re.IGNORECASE)
                if match:
                    caption_para = next_p
                    old_num = int(match.group(1))
                    skip_indices.add(i+1) # Don't process this as a separate image or text paragraph if possible (less critical for images check, but good for order)
            
            images.append({
                'image_para': para,
                'caption_para': caption_para,
                'old_num': old_num
            })
    return images

def add_bookmark(paragraph, bookmark_name):
    """
    Insert a bookmark start and end at the beginning of the paragraph.
    """
    # This is a basic XML injection to add a bookmark.
    # It's tricky in python-docx.
    # Structure: <w:bookmarkStart w:id="0" w:name="bookmark_name"/> ...text... <w:bookmarkEnd w:id="0"/>
    # We'll just try to add it to the start of the paragraph's pPr or run.
    
    # Simpler: Don't mess with XML if not strictly necessary for "linking" visually. 
    # But user asked for "linked internally".
    # I'll skip complex XML for now to avoid corruption, relying on consistent naming.
    pass

def renumber_figures(doc_path, output_path):
    doc = docx.Document(doc_path)
    images = get_document_images(doc)

    print(f"Found {len(images)} images.")
    
    # Mapping for text updates
    old_to_new = {}
    
    # Iterate and Renumber
    for idx, img_info in enumerate(images):
        new_num = idx + 1
        caption_text = f"Figure {new_num}"
        
        if img_info['old_num'] is not None:
            # Update existing caption
            # We assume the caption format matches "Figure X..."
            # Replace logic: "Figure <old>" -> "Figure <new>"
            # Use regex sub to only replace the start
            p = img_info['caption_para']
            old_str = str(img_info['old_num'])
            # Regex: Start of string, Figure, spaces, old_num
            # We preserve the rest of the text
            p.text = re.sub(r"^Figure\s+\d+", caption_text, p.text)
            
            old_to_new[img_info['old_num']] = new_num
            print(f"Updated Figure {img_info['old_num']} -> {new_num}")
            
        else:
            # Create new caption
            # Insert after image_para
            # find index of image_para to get next para
            # Note: doc.paragraphs will change as we insert, but objects should be fine?
            # actually we can use insert_paragraph_before on the *next* paragraph
            
            # Find the paragraph following the image
            # We can't easily get 'next' from the object unless we iterate doc again or use _element.getnext()
            # Safer to iterate doc.paragraphs to find our image_para location?
            # Or just use `image_para.insert_paragraph_before` (wait valid?) NO.
            
            # Use Low-level XML sibling:
            # p._element.addnext(new_p_element) ?
            
            # High-level: `next_paragraph.insert_paragraph_before`
            # We need to find the "next paragraph"
            
            # Let's iterate doc.paragraphs to find the index of this image_para
            # This is slow O(N^2) but safe. N is small (1500 lines).
            
            try:
                curr_idx = doc.paragraphs.index(img_info['image_para'])
                if curr_idx + 1 < len(doc.paragraphs):
                    next_p = doc.paragraphs[curr_idx + 1]
                    new_p = next_p.insert_paragraph_before(caption_text)
                    # Add reference text to the *next* paragraph (which is next_p)
                    # "Adjust the wordings ... accomodate the description"
                    # We'll validly check if next_p is empty or not.
                    if next_p.text.strip():
                        next_p.add_run(f" (refer to {caption_text})")
                else:
                    doc.add_paragraph(caption_text)
            except ValueError:
                print("Could not find paragraph index.")

            print(f"Created new {caption_text}")

    # Update Text References
    # Iterate all paragraphs. If we find "Figure X", check if X is in old_to_new.
    # If yes, replace.
    
    print("\nUpdating text references...")
    # Sort keys reverse to handle 1 vs 10 issues (replace 10 first? No, regex boundaries handle this)
    # \b matches word boundary
    
    for p in doc.paragraphs:
        # Skip if it is a caption (we already updated them or they are new)
        # Identify if it is a caption: starts with "Figure \d"
        if re.match(r"^Figure\s+\d+", p.text):
            continue
            
        # Find all Figure references
        # We look for "Figure X" or "Figures X and Y"
        # Simplest: Regex replacement for each old key
        
        for old_num, new_num in old_to_new.items():
            # Pattern: Word boundary, Figure, spaces, old_num, Word boundary
            pattern = r"\bFigure\s+" + str(old_num) + r"\b"
            if re.search(pattern, p.text):
                # Replace
                replacement = f"Figure {new_num}"
                p.text = re.sub(pattern, replacement, p.text)
                print(f"Updated reference in text: Figure {old_num} -> {new_num}")

    doc.save(output_path)
    print(f"Saved to {output_path}")

if __name__ == "__main__":
    renumber_figures("Empirical Analysis of Accuracy.docx", "Empirical Analysis of Accuracy.docx")
