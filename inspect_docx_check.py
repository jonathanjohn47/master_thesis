
import docx
from docx.shared import Pt

def inspect_docx(file_path):
    doc = docx.Document(file_path)
    
    fonts = set()
    font_sizes = set()
    
    # Check paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts.add(run.font.name)
            if run.font.size:
                font_sizes.add(run.font.size.pt if run.font.size else "Default")
                
    # Check styles (defaults)
    styles = doc.styles
    default_font = styles['Normal'].font.name
    if default_font:
        fonts.add(f"Style:Normal ({default_font})")
        
    print(f"Detected Fonts: {fonts}")
    print(f"Detected Font Sizes: {font_sizes}")
    
    # Word count from doc properties (if available) or manual count
    # docx doesn't give page count easily, but we can count words
    text = [p.text for p in doc.paragraphs]
    full_text = " ".join(text)
    word_count = len(full_text.split())
    print(f"Word Count: {word_count}")
    
    # Check margins (from sections)
    for i, section in enumerate(doc.sections):
        print(f"Section {i+1} Margins:")
        print(f"  Top: {section.top_margin.cm if section.top_margin else 'Default'} cm")
        print(f"  Bottom: {section.bottom_margin.cm if section.bottom_margin else 'Default'} cm")
        print(f"  Left: {section.left_margin.cm if section.left_margin else 'Default'} cm")
        print(f"  Right: {section.right_margin.cm if section.right_margin else 'Default'} cm")

if __name__ == "__main__":
    inspect_docx("Empirical Analysis of Accuracy.docx")
