
from docx import Document

def update_conclusion_in_master(docx_path):
    print(f"Updating conclusion in {docx_path}...")
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error opening file: {e}")
        return
    
    # Refined Conclusion Text
    conclusion_text = [
        ("6. Conclusion", "Heading 1"),
        ("6.1 Summary of Findings", "Heading 2"),
        ("This research empirically evaluated the trade-offs between utility and privacy in federated recommendation systems using the MovieLens 100K dataset. The comprehensive experimental analysis yields three primary conclusions regarding the deployability of private federated learning on mobile devices.", "Normal"),
        ("First, differential privacy proves to be a viable protection mechanism that does not prohibitively sacrifice utility. The results demonstrate that a privacy budget in the range of ε ∈ [4, 8] offers an optimal trade-off, providing rigorous formal privacy guarantees while maintaining recommendation accuracy comparable to non-private baselines. Stricter privacy settings (lower ε) result in a sharp utility drop, identifying this range as the critical operating point for practical deployment.", "Normal"),
        ("Second, the vulnerability analysis reveals that the sparsity of recommendation data acts as a natural privacy defense. Model Inversion Attacks (MIA) failed completely across all experimental configurations, achieving a 0% success rate. Membership Inference Attacks similarly showed limited effectiveness even without differential privacy, suggesting that gradient updates in matrix factorization models leak significantly less sensitive information than those in dense neural networks for image or text classification.", "Normal"),
        ("Third, the system demonstrated unexpected robustness to data heterogeneity. Contrasting with common challenges in federated learning, variations in the non-IID degree (controlled by the Dirichlet parameter α) had negligible impact on model convergence and final accuracy. This indicates that standard federated averaging is sufficiently robust for collaborative filtering tasks even when user data distributions are highly skewed.", "Normal"),
        ("6.2 Limitations and Future Directions", "Heading 2"),
        ("While privacy and heterogeneity concerns were effectively mitigated, the study identified that the primary performance bottleneck is the limited communication budget. The gap between federated and centralized accuracy is driven more by the constrained number of training rounds than by privacy noise or data distribution. Consequently, future work should prioritize communication-efficient optimization strategies, such as gradient compression and quantization, to enable longer training durations on resource-constrained mobile devices.", "Normal"),
        ("Additionally, while the current attack models were ineffective, future research should explore more sophisticated threats tailored to sparse data, such as property inference attacks, to ensure comprehensive security in production environments.", "Normal")
    ]
    
    # Locate existing Conclusion
    start_index = -1
    end_index = -1
    
    # Precise iteration to find 'Conclusion' heading
    for i, p in enumerate(doc.paragraphs):
        text_strip = p.text.strip()
        if ("Conclusion" == text_strip or "6. Conclusion" == text_strip) and p.style.name.startswith("Heading"):
            print(f"Found start at paragraph {i}: '{text_strip}'")
            start_index = i
            # Find end (next Heading 1 or end of doc)
            for j in range(i+1, len(doc.paragraphs)):
                if doc.paragraphs[j].style.name.startswith("Heading 1"):
                    end_index = j
                    print(f"Found end at paragraph {j}: '{doc.paragraphs[j].text}'")
                    break
            if end_index == -1:
                end_index = len(doc.paragraphs)
                print(f"End at paragraph {end_index} (End of doc)")
            break
            
    if start_index != -1:
        print(f"Applying update to range {start_index} - {end_index}")
        
        anchor = doc.paragraphs[start_index]
        
        # Insert New Content BEFORE the anchor
        for text, style in conclusion_text:
            new_p = anchor.insert_paragraph_before(text, style=style)
            
        # Delete Old Content (from original start_index to end_index)
        # We collected indices before insertion, but since we inserted BEFORE start_index,
        # the indices of the old paragraphs *should* remain valid RELATIVE to the doc structure 
        # IF we had references to paragraph objects.
        # But wait, insert_paragraph_before adds to the xml tree.
        # To be safe, we should collect the objects to delete upfront.
        
        # BUT I didn't collect them upfront in the loop above.
        # Let's recollect them based on identity if possible, OR
        # Since we inserted N new paragraphs before index X, the old paragraph at X is now at X+N.
        # Actually, let's just use the fact that we have the 'anchor' object (the old Heading).
        # We can iterate from anchor onwards until we hit the next Heading 1.
        
        pass
        # Re-logic for robust deletion:
        # We want to delete 'anchor' and everything following it until the next Heading 1.
        
        to_delete = []
        curr = anchor
        while curr:
             to_delete.append(curr)
             # Move to next
             # python-docx doesn't imply a linked list structure for paragraphs easily.
             # We rely on the valid iteration.
             
             # Better: use the original list iteration.
             
    else:
        print("Could not find Heading 'Conclusion'. trying looser search...")
        for i, p in enumerate(doc.paragraphs):
             if "Conclusion" in p.text and len(p.text) < 20:
                  print(f"Potential match at {i}: {p.text}")

    # Re-open to get clean indices because insertion shifted things? 
    # No, let's do this:
    # 1. Open doc.
    # 2. Find range [start, end].
    # 3. Collect paragraph objects in that range into a list `paragraphs_to_remove`.
    # 4. Insert new paragraphs before `paragraphs_to_remove[0]`.
    # 5. Remove `paragraphs_to_remove`.
    
    doc = Document(docx_path) # Fresh
    
    p_to_remove = []
    in_section = False
    
    for p in doc.paragraphs:
        if ("Conclusion" == p.text.strip() or "6. Conclusion" == p.text.strip()) and p.style.name.startswith("Heading 1"):
            in_section = True
        
        if in_section:
            # Check if we hit the NEXT Heading 1 (but not the start one)
            if p.style.name.startswith("Heading 1") and p != doc.paragraphs[start_index] if start_index != -1 else False: 
                 # Wait, logic is checking object identity? No.
                 # Let's use text/style.
                 if p.style.name.startswith("Heading 1") and ("Conclusion" not in p.text and "6. Conclusion" not in p.text):
                     break
            p_to_remove.append(p)

    # Note: The loop logic above is slightly flawed because it doesn't stop immediately on the next heading.
    # Revised logic:
    p_to_remove = []
    found_start = False
    
    for p in doc.paragraphs:
        is_heading1 = p.style.name.startswith("Heading 1")
        
        if not found_start:
            if is_heading1 and ("Conclusion" in p.text):
                found_start = True
                p_to_remove.append(p)
        else:
            if is_heading1:
                break # Found next section
            p_to_remove.append(p)
            
    if not p_to_remove:
        print("Still could not find Conclusion section to remove.")
        return

    print(f"Removing {len(p_to_remove)} paragraphs.")
    
    # Insert before the first one
    first_p = p_to_remove[0]
    for text, style in conclusion_text:
        first_p.insert_paragraph_before(text, style=style)
        
    # Remove old ones
    count = 0 
    for p in p_to_remove:
        try:
            p._element.getparent().remove(p._element)
            count += 1
        except Exception as e:
            print(f"Error removing paragraph: {e}")
            
    doc.save(docx_path)
    print(f"Updated {docx_path} successfully. Removed {count} old paragraphs.")

if __name__ == "__main__":
    update_conclusion_in_master("Empirical Analysis of Accuracy.docx")
