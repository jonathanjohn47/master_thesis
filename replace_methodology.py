
import pypdf
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def extract_text_from_pdf(pdf_path):
    try:
        reader = pypdf.PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def replace_methodology(docx_path, pdf_text, dry_run=False):
    try:
        doc = Document(docx_path)
        
        # Locate Methodology Section
        start_index = -1
        end_index = -1
        
        for i, p in enumerate(doc.paragraphs):
            if "Methodology" in p.text and p.style.name.startswith('Heading'):
                start_index = i
            # Find the start of the next major section (e.g. 4. Experimental Setup)
            # We assume it follows Methodology and starts with a digit but is NOT "3.X"
            if start_index != -1 and i > start_index and p.style.name.startswith('Heading'):
                 text = p.text.strip()
                 if text and text[0].isdigit() and not text.startswith("3."):
                     end_index = i
                     break
        
        if start_index == -1:
            print("Could not find 'Methodology' section.")
            return
        
        if end_index == -1:
            print("Could not find the end of 'Methodology' section (start of next chapter). Assuming end of document?")
            # If no next chapter is found, maybe use end of doc? 
            # But let's be safe and ask for manual check if logic fails
            # However, from inspection we know "Experimental Setup" is at 149.
            # So let's fall back to a reasonable guess if needed, or just fail.
            # But wait, inspection showed "Experimental Setup" is present.
            pass

        print(f"Replacing range: {start_index} to {end_index}")
        
        if dry_run:
            print("Dry run complete. No changes made.")
            return

        # Delete paragraphs in reverse order to maintain indices of earlier paragraphs?
        # Actually, python list manipulation: slicing is better but we can't slice doc.paragraphs directly to delete.
        # We have to remove elements from the XML.
        
        # Strategy: 
        # 1. Collect paragraphs to remove.
        # 2. Insert new paragraphs BEFORE the start_index (or at start_index).
        # 3. Remove old paragraphs.
        
        # But `doc.paragraphs` is a list of proxies.
        # To delete, method: p._element.getparent().remove(p._element)
        
        pars_to_remove = doc.paragraphs[start_index+1:end_index]
        
        # Logic: formatting PDF text
        lines = pdf_text.splitlines()
        
        # We need to insert AFTER the paragraph at start_index.
        # But docx doesn't have 'insert_paragraph_after'.
        # We can insert_paragraph_before the paragraph at start_index+1? 
        # But we are about to delete it.
        # So we should insert BEFORE doc.paragraphs[start_index+1] (which is the first one to be removed).
        
        # Wait, if we use insert_paragraph_before on a paragraph that we later remove, the inserted paragraphs might be removed too?
        # No, removing a paragraph element removes THAT element. The implementation of insert_paragraph_before usually adds a sibling element.
        # So it should be fine.
        
        # Let's find the reference paragraph.
        if start_index + 1 < len(doc.paragraphs):
            ref_p = doc.paragraphs[start_index + 1]
        else:
            # If methodology is the last thing, we append?
            # Unlikely given we found an end_index.
            ref_p = None # This simple script might crash if structure is weird.
        
        # Remove the first line of PDF if it matches the title "Methodology"
        if lines and "Methodology" in lines[0]:
            lines = lines[1:]

        for line in lines:
            line = line.strip()
            if not line: continue
            
            # Logic for style
            style = 'Normal'
            if line.startswith("3."):
                # Check depth
                dots = line.count('.')
                # 3.1 -> 1 dot? No 3.1 has 1 dot.
                # 3.1.1 -> 2 dots.
                # But wait, python count('.')
                # "3.1 Research Design" -> 1 dot.
                # "3.7.1" -> 2 dots.
                if dots == 1: style = 'Heading 2' 
                elif dots == 2: style = 'Heading 3'
            
            # Special case for Table and Figure captions if detected?
            # The PDF text has "Table 1: ..." and "Figure 1: ..."
            if line.startswith("Table") or line.startswith("Figure"):
                style = 'Caption'
            
            # Font handling? We rely on styles.
            
            if ref_p:
                new_p = ref_p.insert_paragraph_before(line, style=style)
            else:
                 pass # Should append if no ref_p
            
        # Now remove the old paragraphs
        for p in pars_to_remove:
            # Safely remove
            p._element.getparent().remove(p._element)

        doc.save('Empirical Analysis of Accuracy_Updated.docx')
        print("Document saved to Empirical Analysis of Accuracy_Updated.docx")

    except Exception as e:
        print(f"Error processing DOCX: {e}")

if __name__ == "__main__":
    pdf_path = "Methodology_Section.pdf"
    docx_path = "Empirical Analysis of Accuracy.docx"
    
    pdf_text = extract_text_from_pdf(pdf_path)
    if pdf_text:
        replace_methodology(docx_path, pdf_text)
