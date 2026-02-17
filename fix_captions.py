
import docx
import re
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def fix_captions(doc_path, output_path):
    doc = docx.Document(doc_path)
    
    # Step 1: Clean up dragging captions at the end
    # We suspect the last ~30 paragraphs might contain our accidentally added "Figure X"
    # We iterate backwards
    print("Cleaning up misplaced captions at the end...")
    paras = doc.paragraphs
    removed_count = 0
    for i in range(len(paras) - 1, len(paras) - 50, -1):
        if i < 0: break
        p = paras[i]
        # Match exactly "Figure \d+" with no punctuation (our script generated "Figure X")
        if re.match(r"^Figure\s+\d+$", p.text.strip()):
            print(f"Removing misplaced: '{p.text}'")
            delete_paragraph(p)
            removed_count += 1
    
    print(f"Removed {removed_count} misplaced captions.")
    
    # Step 2: Proper Captioning
    # We need to identify images again.
    # We use list(doc.paragraphs) to get a snapshot of paragraphs to iterate
    # But we need to be careful: if we removed paragraphs, the list might be stale?
    # No, we removed from end. The top part is stable.
    # We will re-read doc.paragraphs to be safe.
    
    all_paras = list(doc.paragraphs) 
    
    image_count = 0
    
    for i, p in enumerate(all_paras):
        # 1. Check for image
        has_image = False
        if 'w:drawing' in p._element.xml:
            has_image = True
        
        if has_image:
            image_count += 1
            expected_caption = f"Figure {image_count}"
            
            # Check if next paragraph is ALREADY a caption
            # We need to find the current valid next sibling in the live document
            # all_paras[i+1] is the *original* next paragraph.
            # If we haven't inserted anything yet, it is still the next paragraph.
            
            if i + 1 >= len(all_paras):
                # Image is at end of doc?
                # We should append caption.
                doc.add_paragraph(expected_caption)
                print(f"Appended {expected_caption} at end.")
                continue
                
            next_p = all_paras[i+1]
            
            # Check if it looks like a caption (Figure \d+)
            # Note: We already renamed existing captions to "Figure X"
            # So "Figure 8" should be there.
            
            if re.match(r"^Figure\s+\d+", next_p.text.strip()):
                # It is a caption.
                # Check if number matches.
                # If we have "Figure 8" and we expect "Figure 8", good.
                # If we have "Figure 8" and we expect something else?
                # The count should align if our first pass was sequential.
                # First pass said "Found 27 images".
                pass
            else:
                # No caption found. This is where we failed last time.
                print(f"Inserting {expected_caption} after image {image_count}")
                next_p.insert_paragraph_before(expected_caption)
                
                # Add text reference to next_p (which is now after the caption)
                # Adjust wording: " (refer to Figure X)"
                # Skip if next_p contains an image (don't add text to an image paragraph)
                if next_p.text.strip() and 'w:drawing' not in next_p._element.xml:
                     next_p.add_run(f" (refer to {expected_caption})")

    doc.save(output_path)
    print(f"Saved fixed document to {output_path}")

if __name__ == "__main__":
    fix_captions("Empirical Analysis of Accuracy.docx", "Empirical Analysis of Accuracy.docx")
