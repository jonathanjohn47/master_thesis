
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_font(run, font_name="Times New Roman", font_size=11):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

def add_formatted_paragraph(doc, text, style='BodyText', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False):
    # If style is a heading, use the docx style
    if 'Heading' in style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    
    p.alignment = alignment
    
    # Simple regex parser for tags: <b>, <i>, <sub>, <sup>, <font>
    # We will split by tags and process chunks
    # Note: this is a simple parser and assumes non-nested tags for simplicity 
    # (though the sources don't seem to have complex nesting)
    
    # Regex to capture tags and content
    # Matches <tag>content</tag> or just text
    # But wait, splitting is safer.
    
    parts = re.split(r'(<[^>]+>)', text)
    
    current_bold = bold
    current_italic = italic
    current_sub = False
    current_sup = False
    
    for part in parts:
        if not part:
            continue
            
        if part == '<b>':
            current_bold = True
        elif part == '</b>':
            current_bold = False
        elif part == '<i>':
            current_italic = True
        elif part == '</i>':
            current_italic = False
        elif part == '<sub>':
            current_sub = True
        elif part == '</sub>':
            current_sub = False
        elif part == '<sup>':
            current_sup = True
        elif part == '</sup>':
            current_sup = False
        elif part.startswith('<font') or part == '</font>':
            # Ignore font tags for now, usually just for Courier in the source
            pass
        elif part == '<br/>':
            p.add_run().add_break()
        else:
            # Regular text
            run = p.add_run(part)
            run.bold = current_bold
            run.italic = current_italic
            if current_sub:
                run.font.subscript = True
            if current_sup:
                run.font.superscript = True
            
            # Set font style
            if 'Heading' in style:
                pass # Use style default
            elif 'Equation' in style:
                set_font(run, "Cambria Math", 11)
            elif 'Code' in style:
                set_font(run, "Courier New", 9)
            elif 'Caption' in style:
                 set_font(run, "Times New Roman", 10)
                 run.bold = True
            else:
                set_font(run, "Times New Roman", 11)

    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        set_font(run, "Times New Roman", 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            p = row_cells[i].paragraphs[0]
            # Process formatting in cell data if any
            # For simplicity, strip tags or just add as text
            # The source has minimal formatting in tables, mostly text
            # But let's use our formatter manually
            
            # Start formatting
            parts = re.split(r'(<[^>]+>)', str(cell_data))
            c_bold = False; c_italic = False; c_sub = False; c_sup = False
            for part in parts:
                if part == '<b>': c_bold = True
                elif part == '</b>': c_bold = False
                elif part == '<i>': c_italic = True
                elif part == '</i>': c_italic = False
                elif part == '<sub>': c_sub = True
                elif part == '</sub>': c_sub = False
                elif part == '<sup>': c_sup = True
                elif part == '</sup>': c_sup = False
                elif part.startswith('<') and part.endswith('>'): pass
                else:
                    if not part: continue
                    run = p.add_run(part)
                    run.bold = c_bold; run.italic = c_italic
                    if c_sub: run.font.subscript = True
                    if c_sup: run.font.superscript = True
                    set_font(run, "Times New Roman", 10)
            
            # Align center usually, except maybe first col
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer

def add_figure(doc, filename, caption):
    fig_path = os.path.join(os.getcwd(), 'figures', filename)
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Caption
        p = doc.add_paragraph(style='Caption')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        set_font(run, "Times New Roman", 10)
    else:
        p = doc.add_paragraph(f"[Image not found: {filename}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    doc = Document()
    
    # 3. METHODOLOGY
    doc.add_heading('3. Methodology', 0)
    
    add_formatted_paragraph(doc, "3.1 Research Design", "Heading 2")
    add_formatted_paragraph(doc, "This study employs an empirical, quantitative research design to investigate the accuracy-privacy trade-offs in federated learning for mobile movie recommendation systems. The research is structured around three research questions that examine (RQ1) the impact of differential privacy budgets on recommendation accuracy, (RQ2) the effectiveness of privacy attacks under varying DP configurations, and (RQ3) the influence of data heterogeneity on federated learning performance.")
    add_formatted_paragraph(doc, "The experimental framework follows a controlled variable approach: for RQ1, the privacy budget (\u03b5) is varied while holding data distribution constant; for RQ2, privacy attacks are evaluated across all DP configurations; and for RQ3, the Dirichlet concentration parameter (\u03b1) is varied while disabling differential privacy. Each configuration is repeated with three random seeds (42, 123, 456) to ensure statistical reliability and to report mean \u00b1 standard deviation of results.")
    
    add_formatted_paragraph(doc, "3.2 Dataset", "Heading 2")
    add_formatted_paragraph(doc, "We use the MovieLens 100K dataset (Harper & Konstan, 2015), a widely adopted benchmark for recommendation system research. The dataset contains 100,000 ratings from 943 users across 1,682 movies, with ratings on a 1\u20135 integer scale. The dataset exhibits 93.7% sparsity, which is representative of real-world recommendation scenarios. Ratings are used in their original continuous form (1\u20135 scale) rather than binarized, enabling regression-based evaluation that better captures the nuances of user preference prediction.")
    add_formatted_paragraph(doc, "The dataset is split into 80% training (80,000 interactions) and 20% test (20,000 interactions) sets using a random permutation with a fixed seed for reproducibility. For the federated learning experiments, the training data is further partitioned among 100 simulated clients using a Dirichlet distribution to create non-IID (non-independent and identically distributed) data splits, which models the natural heterogeneity of user preferences across mobile devices.")
    
    add_formatted_paragraph(doc, "3.3 Model Architecture", "Heading 2")
    add_formatted_paragraph(doc, "The recommendation model is based on Matrix Factorization (MF), a proven approach for collaborative filtering that decomposes the user-item interaction matrix into low-dimensional latent factor representations. The model consists of two embedding layers:")
    add_formatted_paragraph(doc, "\u2022 <b>User Embedding</b>: Maps each of the 943 users to a 64-dimensional latent vector.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "\u2022 <b>Item Embedding</b>: Maps each of the 1,682 items to a 64-dimensional latent vector.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "The predicted rating for a user-item pair (u, i) is computed as the dot product of the corresponding user and item embedding vectors:")
    
    # Equation
    add_formatted_paragraph(doc, "\u0177(u, i) = e<sub>u</sub><sup>T</sup> \u00b7 e<sub>i</sub>", "Equation", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_formatted_paragraph(doc, "Embeddings are initialized with Gaussian noise (std = 0.1) to ensure adequate gradient flow during early training. The embedding dimension of 64 was selected to provide sufficient representational capacity while remaining deployable on mobile devices. The total model size is (943 + 1,682) \u00d7 64 = 167,936 parameters, which is suitable for on-device training on resource-constrained mobile platforms.")
    
    add_formatted_paragraph(doc, "3.4 Federated Learning Protocol", "Heading 2")
    add_formatted_paragraph(doc, "The federated learning protocol follows the Federated Averaging (FedAvg) algorithm (McMahan et al., 2017). The protocol operates over T = 10 communication rounds between a central server and K = 100 simulated clients. Each round proceeds as follows:")
    add_formatted_paragraph(doc, "<b>Step 1 \u2013 Global Distribution:</b> The server broadcasts the current global model parameters to all participating clients.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "<b>Step 2 \u2013 Local Training:</b> Each client trains the received model on their local data for E = 3 local epochs using Stochastic Gradient Descent (SGD) with a learning rate of 0.01, batch size of 32, and L2 regularization (weight decay = 10<sup>\u20135</sup>). The loss function is Mean Squared Error (MSE) for regression-based rating prediction.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "<b>Step 3 \u2013 Parameter Upload:</b> Each client sends its updated model parameters back to the server.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "<b>Step 4 \u2013 Aggregation:</b> The server performs weighted averaging of client parameters, where weights are proportional to each client\u2019s local dataset size:", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_formatted_paragraph(doc, "w<sub>global</sub> = \u03a3<sub>k=1</sub><sup>K</sup> (n<sub>k</sub> / n<sub>total</sub>) \u00b7 w<sub>k</sub>", "Equation", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_formatted_paragraph(doc, "where n<sub>k</sub> is the number of training samples on client k and n<sub>total</sub> is the total number of samples across all clients.")
    
    add_formatted_paragraph(doc, "3.5 Non-IID Data Partitioning", "Heading 2")
    add_formatted_paragraph(doc, "To simulate realistic mobile data distributions, we partition the training data among clients using a Dirichlet distribution with concentration parameter \u03b1. For each user in the dataset, a probability vector is sampled from Dir(\u03b1, \u2026, \u03b1) with K = 100 components, determining the proportion of that user\u2019s interactions assigned to each client. The parameter \u03b1 controls the degree of heterogeneity:")
    add_formatted_paragraph(doc, "\u2022 <b>\u03b1 = 0.1 (Highly Non-IID)</b>: Most of a user\u2019s data is concentrated on a single client, creating highly skewed distributions.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "\u2022 <b>\u03b1 = 0.5 (Moderately Non-IID)</b>: An intermediate distribution representing realistic mobile scenarios where users have some overlap but distinct preferences.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "\u2022 <b>\u03b1 = 1.0 (Mildly Non-IID)</b>: Data is more evenly distributed across clients, approaching but not reaching the IID case.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_formatted_paragraph(doc, "3.6 Differential Privacy Mechanism", "Heading 2")
    add_formatted_paragraph(doc, "We implement Differentially Private Stochastic Gradient Descent (DP-SGD) (Abadi et al., 2016) to provide formal privacy guarantees during local training. The mechanism consists of two steps applied to each gradient computation during training:")
    add_formatted_paragraph(doc, "<b>Gradient Clipping:</b> Per-sample gradients are clipped to a maximum L2 norm of C = 1.0. This bounds the sensitivity of each individual training sample:", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "g\u0302 = g \u00b7 min(1, C / \u2016g\u2016<sub>2</sub>)", "Equation", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "<b>Gaussian Noise Addition:</b> After clipping, calibrated Gaussian noise is added to the gradients:", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_formatted_paragraph(doc, "g\u0303 = g\u0302 + \U0001d4a9(0, \u03c3\u00b2 \u00b7 C\u00b2 \u00b7 I)", "Equation", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "where \u03c3 is the noise multiplier that determines the privacy-utility trade-off. The noise multiplier is calibrated using the R\u00e9nyi Differential Privacy (RDP) accountant (Mironov, 2017) to achieve a target (\u03b5, \u03b4)-DP guarantee with \u03b4 = 10<sup>\u20135</sup>. The RDP composition theorem tracks the cumulative privacy loss across multiple rounds of training. We evaluate five privacy budgets:")
    
    add_table(doc, 
        ['Target \u03b5', '\u03c3 (Noise Multiplier)', 'Achieved \u03b5', 'Privacy Level'],
        [
            ['\u221e', '0.00', '\u221e', 'No Privacy'],
            ['8', '11.03', '7.91', 'Relaxed'],
            ['4', '20.40', '4.02', 'Moderate'],
            ['2', '40.70', '1.94', 'Strong'],
            ['1', '75.06', '1.03', 'Very Strong'],
        ]
    )
    
    add_formatted_paragraph(doc, "3.7 Privacy Attack Evaluation", "Heading 2")
    add_formatted_paragraph(doc, "3.7.1 Membership Inference Attack (MIA)", "Heading 3")
    add_formatted_paragraph(doc, "We implement a shadow-model based Membership Inference Attack (Shokri et al., 2017) to evaluate whether an adversary can determine if a specific data sample was used in the training set.")
    
    add_formatted_paragraph(doc, "3.7.2 Model Inversion Attack", "Heading 3")
    add_formatted_paragraph(doc, "We implement a model inversion attack that attempts to reconstruct a user\u2019s top-K preferred items from the trained model.")
    
    add_formatted_paragraph(doc, "3.8 Evaluation Metrics", "Heading 2")
    add_formatted_paragraph(doc, "We evaluate recommendation quality using NDCG@10, Hit@10, Precision@10, and Recall@10, alongside regression metrics MSE and MAE.")
    add_formatted_paragraph(doc, "NDCG@K = DCG@K / IDCG@K", "Equation", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 4. EXPERIMENTAL SETUP
    doc.add_page_break()
    doc.add_heading('4. Experimental Setup', 0)
    
    add_formatted_paragraph(doc, "4.1 Overview", "Heading 2")
    add_formatted_paragraph(doc, "This chapter details the complete experimental setup used to evaluate the accuracy\u2013privacy trade-offs in federated learning for mobile movie recommendation systems. The experiments are designed to answer three research questions: (RQ1) the impact of differential privacy budgets on recommendation accuracy, (RQ2) the effectiveness of privacy attacks under varying DP configurations, and (RQ3) the influence of data heterogeneity on federated learning performance.")
    
    add_formatted_paragraph(doc, "4.2 Hardware and Software Environment", "Heading 2")
    add_formatted_paragraph(doc, "All experiments are executed on a single workstation equipped with a multi-core CPU. No GPU acceleration is employed; the relatively compact model size (approximately 168K parameters) makes CPU-only training practical.")
    
    add_table(doc, 
        ["Component", "Library / Framework", "Version", "Purpose"],
        [
            ["Deep Learning", "PyTorch", "\u2265 2.0.0", "Model training"],
            ["Numerical", "NumPy", "\u2265 1.24.0", "Data partitioning"],
            ["Data Processing", "Pandas", "\u2265 2.0.0", "Result aggregation"],
            ["ML Utilities", "Scikit-learn", "\u2265 1.3.0", "Attack classifier"],
            ["Visualization", "Matplotlib / Seaborn", "\u2265 3.7", "Figures"],
        ]
    )
    
    add_formatted_paragraph(doc, "4.3 Dataset Preparation", "Heading 2")
    add_formatted_paragraph(doc, "The MovieLens 100K dataset is used. It is split into 80% training (80,000 interactions) and 20% test (20,000 interactions) sets.")
    add_formatted_paragraph(doc, "To simulate realistic mobile data distributions, the training data is partitioned among K=100 clients using a Dirichlet distribution with concentration parameter \u03b1.")
    
    add_formatted_paragraph(doc, "Table 4: Dirichlet Concentration Parameter Configurations", "Caption")
    add_table(doc,
        ["\u03b1", "Heterogeneity Level", "Typical Client Profile"],
        [
            ["0.1", "Highly Non-IID", "Data from 1-2 users"],
            ["0.5", "Moderately Non-IID", "Data from several users"],
            ["1.0", "Mildly Non-IID", "Approaching IID"],
        ]
    )

    add_formatted_paragraph(doc, "4.4 Federated Learning Simulation", "Heading 2")
    add_formatted_paragraph(doc, "The experiments are conducted using an in-memory simulation that replicates the Federated Averaging (FedAvg) protocol. In each round, every client receives the global model, performs local training, and returns updates.")
    
    add_formatted_paragraph(doc, "4.7 Experiment Configurations", "Heading 2")
    add_formatted_paragraph(doc, "The experimental design consists of four experiment groups, totaling 30 individual experiment runs. Each configuration is repeated with three random seeds (42, 123, 456).")
    
    add_table(doc, 
        ["Experiment", "Clients", "Rounds", "\u03b5", "\u03b1", "Seeds"],
        [
            ["Centralized", "N/A", "50", "N/A", "N/A", "1"],
            ["DP Sweep (RQ1)", "100", "10", "{\u221e,8,4,2,1}", "0.5", "3"],
            ["Attacks (RQ2)", "100", "10", "{\u221e,8,4,2,1}", "0.5", "1"],
            ["Heterogeneity (RQ3)", "100", "10", "\u221e", "{0.1,0.5,1.0}", "3"],
        ]
    )

    # 5. RESULTS
    doc.add_page_break()
    doc.add_heading('5. Results', 0)
    
    add_formatted_paragraph(doc, "5.1 Overview", "Heading 2")
    add_formatted_paragraph(doc, "This chapter presents the results of the 30 experiments conducted to evaluate the accuracy\u2013privacy trade-offs. The results are organized around the three research questions.")

    add_formatted_paragraph(doc, "5.2 Centralized Baseline", "Heading 2")
    add_formatted_paragraph(doc, "The centralized model serves as the upper bound for recommendation accuracy.")
    
    add_table(doc,
        ["Metric", "Value"],
        [
            ["NDCG@10", "0.2250"],
            ["Hit@10", "0.3800"],
            ["MSE", "2.045"],
            ["MAE", "1.106"],
        ]
    )
    
    add_formatted_paragraph(doc, "5.3 RQ1: Impact of Differential Privacy", "Heading 2")
    add_formatted_paragraph(doc, "Table 9 presents the recommendation quality metrics across all five DP budget configurations.")
    
    add_table(doc,
        ["\u03b5", "NDCG@10", "Hit@10"],
        [
            ["\u221e (No DP)", "0.0539 \u00b1 0.011", "0.0633 \u00b1 0.021"],
            ["8", "0.0534 \u00b1 0.013", "0.0600 \u00b1 0.008"],
            ["4", "0.0479 \u00b1 0.009", "0.0600 \u00b1 0.008"],
            ["2", "0.0456 \u00b1 0.010", "0.0567 \u00b1 0.009"],
            ["1", "0.0467 \u00b1 0.012", "0.0533 \u00b1 0.005"],
        ]
    )
    
    add_figure(doc, "accuracy_vs_epsilon.png", "Figure 1: NDCG@10 and Hit@10 as a function of DP budget (\u03b5).")
    
    add_formatted_paragraph(doc, "Figure 1 illustrates the substantial gap between the centralized baseline and all federated configurations. Even without DP, the federated model achieves only NDCG@10 = 0.054, representing a 76% degradation from the centralized baseline.")

    add_figure(doc, "accuracy_loss_vs_epsilon.png", "Figure 2: Relative accuracy loss (%) vs. DP budget (\u03b5).")
    
    add_formatted_paragraph(doc, "5.3.4 Training Convergence", "Heading 3")
    add_figure(doc, "convergence.png", "Figure 3: Training loss convergence (left) and NDCG@10 convergence (right).")
    
    add_formatted_paragraph(doc, "5.4 RQ2: Effectiveness of Privacy Attacks", "Heading 2")
    add_formatted_paragraph(doc, "The membership inference attack is remarkably ineffective across all configurations, even without differential privacy.")
    
    add_table(doc,
        ["\u03b5", "MIA AUC", "MIA Accuracy"],
        [
            ["\u221e (No DP)", "0.548", "0.545"],
            ["8", "0.534", "0.515"],
            ["4", "0.502", "0.505"],
            ["2", "0.500", "0.500"],
            ["1", "0.486", "0.510"],
        ]
    )
    
    add_figure(doc, "attack_evaluation.png", "Figure 4: Privacy attack effectiveness as a function of DP budget.")
    
    add_formatted_paragraph(doc, "5.5 RQ3: Impact of Data Heterogeneity", "Heading 2")
    add_formatted_paragraph(doc, "The heterogeneity sweep yields a striking result: data heterogeneity has virtually no measurable impact on federated learning performance in this experimental configuration.")
    
    add_table(doc,
        ["\u03b1", "NDCG@10", "Hit@10"],
        [
            ["0.1", "0.0538 \u00b1 0.011", "0.0633 \u00b1 0.021"],
            ["0.5", "0.0539 \u00b1 0.011", "0.0633 \u00b1 0.021"],
            ["1.0", "0.0539 \u00b1 0.011", "0.0633 \u00b1 0.021"],
        ]
    )
    
    add_figure(doc, "accuracy_vs_alpha.png", "Figure 5: NDCG@10 and Hit@10 as a function of data heterogeneity (\u03b1).")

    add_formatted_paragraph(doc, "5.6 Client Data Distribution Analysis", "Heading 2")
    add_figure(doc, "client_distribution.png", "Figure 6: Distribution of training samples per client.")
    add_figure(doc, "aggregation_stats.png", "Figure 7: Aggregation statistics over rounds.")
    
    add_formatted_paragraph(doc, "5.7 Recommendation Metrics Over Training Rounds", "Heading 2")
    add_figure(doc, "recommendation_metrics.png", "Figure 8: Recommendation metrics over training rounds.")

    add_formatted_paragraph(doc, "5.9 Summary of Key Findings", "Heading 2")
    add_formatted_paragraph(doc, "In summary, the results demonstrate that federated learning for mobile recommendation faces a significant accuracy challenge compared to centralized training, but this gap is primarily attributable to the limited communication budget and data fragmentation\u2014not to differential privacy. DP at moderate levels (\u03b5 = 4\u20138) provides robust formal privacy guarantees while imposing only a small additional accuracy penalty.")

    doc.save('Combined_Thesis_Sections.docx')
    print("Document saved to Combined_Thesis_Sections.docx")

if __name__ == "__main__":
    main()
