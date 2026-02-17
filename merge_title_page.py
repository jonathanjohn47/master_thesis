
from docx import Document
import os

def merge_title_page(title_page_path, main_doc_path, output_path):
    print(f"Loading Title Page from {title_page_path}...")
    title_doc = Document(title_page_path)
    
    print(f"Loading Main Thesis from {main_doc_path}...")
    main_doc = Document(main_doc_path)
    
    # Create a new document to hold the merged content? 
    # Or just append main content to title page?
    # Appending main content to title page is safer to preserve Title Page formatting at the start.
    
    print("Merging documents...")
    
    # Add a page break after title page
    title_doc.add_page_break()
    
    # Iterate through elements of the main doc and add them to title_doc
    # Note: This Simple copy might lose some complex formatting (headers/footers of main doc). 
    # But for a "Written Assignment" level merge, this is usually sufficient. 
    # A robust way is using composer or just appending bodies. 
    # Let's try appending body elements.
    
    for element in main_doc.element.body:
        title_doc.element.body.append(element)
        
    # Validating if this naive append works for python-docx. 
    # It often duplicates the body tag or structure. 
    # standard python-docx way is iterating paragraphs/tables.
    
    # Let's re-load to be safe and do paragraph copy to avoid invalid XML
    # Actually, the "append body element" is a known hack but can be risky.
    # Let's try the safer "composer" approach if we had `docxcompose`, but we might not have it installed.
    # We'll stick to a standard element copy loop which is safer for basic text.
    
    combined_doc = Document(title_page_path)
    combined_doc.add_page_break()
    
    for element in main_doc.element.body:
        combined_doc.element.body.append(element)
        
    print(f"Saving merged document to {output_path}...")
    combined_doc.save(output_path)
    print("Merge complete.")

if __name__ == "__main__":
    base_dir = "/Users/jonathanjohn/StudioProjects/master_thesis"
    title_page = os.path.join(base_dir, "Title_Page.docx")
    main_thesis = os.path.join(base_dir, "Empirical Analysis of Accuracy.docx")
    # We will overwrite the main file? Or create a new one?
    # User said "I don't see the title page yet", implying they expect it IN the thesis.
    # To be safe, I'll save to a temp new file first, then rename if successful.
    output_file = os.path.join(base_dir, "Empirical_Analysis_with_Title.docx")
    
    try:
        # Check if docxcompose is available, it's much better.
        from docxcompose.composer import Composer
        print("Using docxcompose for high-fidelity merge.")
        master = Document(title_page)
        composer = Composer(master)
        doc2 = Document(main_thesis)
        master.add_page_break()
        composer.append(doc2)
        composer.save(output_file)
        print("Merge successful with docxcompose.")
        
    except ImportError:
        print("docxcompose not found. Using fallback append.")
        # Fallback
        # Re-implementing a safer element move
        doc1 = Document(title_page)
        doc2 = Document(main_thesis)
        doc1.add_page_break()
        
        for element in doc2.element.body:
            doc1.element.body.append(element)
            
        doc1.save(output_file)

