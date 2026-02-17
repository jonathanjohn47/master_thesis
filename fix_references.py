
import docx
import re

def fix_references(doc_path, output_path):
    doc = docx.Document(doc_path)
    
    # Identify protected paragraphs (Captions)
    # We scanned 27 images.
    # We expect captions at specific places.
    # Heuristic: Paragraphs that start with "Figure \d+" AND are typically adjacent to images
    # OR the ones we know are captions.
    # Better: Scan for images. The very next paragraph (or the one after) checking for "Figure X" matches.
    
    protected_indices = set()
    protected_texts = []
    
    # Map Image Index -> Caption Paragraph Index
    img_idx = 0
    for i, p in enumerate(doc.paragraphs):
        if 'w:drawing' in p._element.xml:
            img_idx += 1
            # Check next few paragraphs for caption
            # We inserted them immediately after.
            # So usually i+1
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i+1]
                if re.match(r"^Figure\s+\d+", next_p.text.strip()):
                    protected_indices.add(i+1)
                    protected_texts.append(next_p.text[:30])
    
    print(f"Protected {len(protected_indices)} caption paragraphs.")
    
    # Define Mapping Rule: Old X -> New (X+7)
    # Range 1 to 8.
    mapping = {x: x + 7 for x in range(1, 9)} # 1..8
    
    # Apply replacements in REVERSE order of keys (8 down to 1)
    # to avoid chaining issues (e.g. 1->8, then 8->15) if we had overlaps in the source side.
    # Here, 8->15. 1->8.
    # If we replace 1->8 first, "Figure 1" becomes "Figure 8".
    # Then if we replace 8->15, that "Figure 8" becomes "Figure 15". WRONG.
    # So we MUST replace 8->15 FIRST. Then 1->8.
    
    sorted_keys = sorted(mapping.keys(), reverse=True)
    
    count = 0
    for i, p in enumerate(doc.paragraphs):
        if i in protected_indices:
            continue
            
        # Optimization: Skip if no "Figure" in text
        if "Figure" not in p.text:
            continue
            
        original_text = p.text
        new_text = original_text
        
        for old_num in sorted_keys:
            new_num = mapping[old_num]
            # Replace complete word "Figure X"
            pattern = r"\bFigure\s+" + str(old_num) + r"\b"
            if re.search(pattern, new_text):
                 replacement = f"Figure {new_num}"
                 new_text = re.sub(pattern, replacement, new_text)
        
        if new_text != original_text:
            p.text = new_text
            count += 1
            print(f"Updated Para {i}: ...{original_text[:40]}... -> ...{new_text[:40]}...")

    doc.save(output_path)
    print(f"Updated {count} text references. Saved to {output_path}")

if __name__ == "__main__":
    fix_references("Empirical Analysis of Accuracy.docx", "Empirical Analysis of Accuracy.docx")
