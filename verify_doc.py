
import docx

def verify(doc_path):
    doc = docx.Document(doc_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    print("\n--- Last 20 Paragraphs ---")
    for i, p in enumerate(doc.paragraphs[-20:]):
        print(f"[{len(doc.paragraphs)-20+i}] {repr(p.text)}")
        
    print("\n--- Captions Found ---")
    count = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.lower().startswith("figure"):
             print(f"[{i}] {text[:50]}")
             count += 1
             
    print(f"Total lines starting with 'Figure': {count}")

if __name__ == "__main__":
    verify("Empirical Analysis of Accuracy.docx")
